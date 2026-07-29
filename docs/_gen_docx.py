# -*- coding: utf-8 -*-
"""
Generates ClientHub_Project_Report.docx — a formal, recruiter-ready MCA project report.
Author of project: Suryakanta Pradhan.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x35, 0x5E)
BLUE = RGBColor(0x2B, 0x6C, 0xB0)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x6B, 0x6B, 0x6B)

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.size = Pt(16)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = BLUE
    r.font.size = Pt(13)
    r.font.bold = True
    return p


def para(text, size=11, color=None, bold=False, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.font.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p


def code_block(lines):
    """Render monospace text inside a shaded single-cell table."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade_cell(cell, 'F3F4F6')
    cell.paragraphs[0].text = ''
    first = True
    for ln in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        r = p.add_run(ln)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    return t


def kv_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].paragraphs[0].text = ''
        r = hdr[i].paragraphs[0].add_run(htext)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], '2B6CB0')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].text = ''
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)
    return t


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

para('ClientHub', size=40, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('AI-Powered Agency Client & Project Management Portal',
     size=15, color=BLUE, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
para('A SaaS-style platform unifying client management, project & task tracking, '
     'invoicing with online payments, real-time collaboration, and AI-generated '
     'progress reporting.', size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)

for _ in range(3):
    doc.add_paragraph()

para('PROJECT REPORT', size=13, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

para('Submitted by', size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('Suryakanta Pradhan', size=18, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('Master of Computer Applications (MCA)', size=11, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('Full-Stack Web Development Project', size=11, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=30)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (static)
# ============================================================
h1('Table of Contents')
toc_items = [
    '1.  Abstract / Overview',
    '2.  Problem Statement',
    '3.  Objectives',
    '4.  Scope of the Project',
    '5.  Complete Feature List',
    '6.  System Architecture',
    '7.  Database Schema (ER Description)',
    '8.  Technology Stack & Justification',
    '9.  Module-wise Explanation',
    '10. The AI Progress Summary Feature (Core Differentiator)',
    '11. Security Considerations',
    '12. Final Project Folder Structure',
    '13. Setup, Seed Data & Demo Credentials',
    '14. Testing, Code Quality & CI/CD',
    '15. Deployment Architecture',
    '16. Screenshots (Placeholders)',
    '17. Challenges Faced & Solutions',
    '18. Future Scope',
    '19. Conclusion',
    '20. Resume Summary',
]
for item in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. ABSTRACT
# ============================================================
h1('1. Abstract / Overview')
para('ClientHub is a production-quality, SaaS-style web application built for freelance '
     'agencies and creative studios to manage their entire client lifecycle in one place. '
     'It brings together three capabilities that companies usually buy as separate tools: '
     '(1) client, project, and task management like a services company needs; (2) a '
     'multi-tenant, role-based dashboard like a product company builds; and (3) invoicing '
     'with online payment collection like a fintech company builds.')
para('The platform serves three distinct user roles — the Agency Owner (Admin), Team Members, '
     'and Clients — each with a tailored, access-controlled experience. Its standout feature '
     "is an AI Progress Summary engine: the system aggregates a project\'s weekly activity "
     '(completed tasks, pending work, hours logged, deadlines, and blockers) and uses the '
     'Groq API to generate a friendly, professional natural-language progress report along '
     'with an automatic risk flag (on-track, at-risk, or delayed). This turns raw project data '
     'into client-ready communication automatically.')

# ============================================================
# 2. PROBLEM STATEMENT
# ============================================================
h1('2. Problem Statement')
para('Freelance agencies and small studios juggle a fragmented toolchain — spreadsheets for '
     'clients, a separate board for tasks, email for status updates, and yet another tool for '
     'invoicing. This fragmentation creates real, recurring problems:')
bullet(' clients have no transparent, self-service view of their project\'s progress and must email for updates.', 'No single source of truth:')
bullet(' writing weekly status updates for every client is repetitive, time-consuming manual work.', 'Manual reporting overhead:')
bullet(' invoicing lives outside the project tool, so billing is disconnected from the work actually done.', 'Disconnected billing:')
bullet(' without a structured view, slipping deadlines are noticed late, after they have already become a problem.', 'Late risk detection:')
para('ClientHub solves these by unifying everything into one role-aware portal and layering an '
     'AI reporting engine on top, so status communication and risk detection become automatic '
     'rather than manual.')

# ============================================================
# 3. OBJECTIVES
# ============================================================
h1('3. Objectives')
objectives = [
    'Build a secure, multi-role platform (Admin, Team Member, Client) with strict role-based access control.',
    'Provide complete client, project, and task management with a drag-and-drop Kanban workflow.',
    'Implement professional invoicing with line items, tax, PDF export, and live online payments via a payment gateway in test mode.',
    'Deliver real-time notifications for task assignments, comments, payments, and status changes.',
    'Give Admins and Clients clear analytics dashboards built on real project and revenue data.',
    'Automatically generate client-ready, AI-written weekly progress summaries with a project risk flag.',
    'Write clean, well-commented, professionally organized code suitable for a public portfolio and recruiter review.',
]
for o in objectives:
    numbered(o)

# ============================================================
# 4. SCOPE
# ============================================================
h1('4. Scope of the Project')
para('The project covers the full stack end-to-end: a React single-page frontend, a Node.js + '
     'Express REST API, a PostgreSQL database accessed through the Prisma ORM, real-time '
     'communication via Socket.io, third-party integrations (Groq AI, a payment gateway, '
     'ImageKit, email), automated background jobs, automated tests, API documentation, and '
     'cloud deployment. The system is designed to be demonstrated live with realistic seeded '
     'demo data and fixed login credentials for each role.')

# ============================================================
# 5. FEATURE LIST
# ============================================================
h1('5. Complete Feature List')

h2('Module 1 — Authentication & Access Control')
for x in ['Email + password signup and login with bcrypt password hashing.',
          'JWT-based stateless session tokens.',
          'Role-based access middleware (admin / team_member / client).',
          'Forgot-password / reset-password via emailed link; reset tokens are SHA-256 hashed and stored in the database (single-use, 1-hour expiry, atomic redemption) so they survive server restarts and work across multiple instances.',
          'Protected frontend routes that redirect based on the user\'s role.',
          'Input validation and rate-limiting on authentication endpoints.']:
    bullet(x)

h2('Module 2 — Client Management')
for x in ['Add / edit / archive / delete client records with logo upload via ImageKit.',
          'Each client record links to a client-role user account for portal login.',
          'Search and filter clients by name, status, or project count.',
          'Client list view plus a detailed client profile showing all their projects.']:
    bullet(x)

h2('Module 3 — Project Management')
for x in ['Create / edit / archive projects with client link, budget, dates, and status.',
          'Project detail page: tasks, assigned team, budget used vs. total, and timeline.',
          'Auto-calculated project progress percentage from completed tasks.']:
    bullet(x)

h2('Module 4 — Task Management (Kanban Board & My Tasks)')
for x in ['Tasks belong to a project with assignee, priority, status, and due date.',
          'Kanban board with four columns: To Do / In Progress / Review / Done.',
          'Drag-and-drop between columns using @dnd-kit.',
          'Comment thread per task, time tracking, and a per-task activity log.',
          'My Tasks view: a cross-project table of every task assigned to the logged-in user, sorted by due date (soonest first, nulls last), with overdue dates highlighted in red and direct links to the parent project.']:
    bullet(x)

h2('Module 5 — Invoicing & Payments (Fintech Layer)')
for x in ['Invoice builder with multiple line items, quantity, rate, auto tax and total.',
          'Auto-suggest invoice line items from logged hours per task.',
          'PDF export of invoices.',
          'Razorpay test-mode integration with a client-facing "Pay Now" flow.',
          'Webhook endpoint with signature verification that auto-updates invoice status.',
          'Payment history log per client with downloadable receipts.']:
    bullet(x)

h2('Module 6 — Real-Time Notifications')
for x in ['Socket.io on backend and frontend for live events.',
          'Events: task assigned, comment added, invoice paid, project status changed.',
          'Notification bell with unread count and a clickable notification list.',
          'Email fallback for high-importance events.']:
    bullet(x)

h2('Module 7 — Analytics Dashboards')
for x in ['Admin: total revenue, pending payments, active vs. completed projects, team workload, top clients.',
          'Client: project progress %, upcoming deadlines, invoice summary, recent activity feed.',
          'Charts built with Recharts.']:
    bullet(x)

h2('Module 8 — AI Progress Summary (Core Differentiator)')
for x in ['Aggregates weekly project activity into a structured prompt.',
          'Calls the Groq API to produce a natural-language summary plus a risk flag.',
          'Stores summaries with history; shows the latest on the client dashboard.',
          'Admin "Generate Summary Now" button plus a weekly automated cron job.',
          'Graceful fallback to a deterministic template summary if the AI call fails.']:
    bullet(x)

h2('Module 9 — User Experience, Theming & Accessibility')
for x in ['Full light/dark theme across every page, following the operating-system preference '
          'on a first visit and remembering the user\'s choice afterwards.',
          'The theme is applied by a tiny inline script before React mounts, so a dark-mode '
          'user never sees a flash of the light background while the page loads.',
          'Responsive layout: the sidebar rail becomes a slide-out drawer below tablet width, '
          'so the application is fully navigable on a phone.',
          'In-house toast notifications (no external dependency) give immediate feedback on '
          'every create, update, archive, download and payment action.',
          'Accessibility: modals trap focus while open, restore focus to whatever opened them, '
          'close on Escape, and lock background scrolling; Kanban cards are reachable by '
          'keyboard and open with Enter or Space; interactive controls carry ARIA labels.',
          'Empty states, loading spinners and a friendly 404 page so no screen is ever blank.']:
    bullet(x)

h2('Module 10 — Performance & Reliability Engineering')
for x in ['Route-level code splitting with React.lazy, plus a dynamic import for the Socket.io '
          'client, so the charting and drag-and-drop libraries download only on the pages that '
          'use them — the initial JavaScript payload dropped from roughly 748 KB to 233 KB.',
          'The client search box is debounced, turning one API request per keystroke into one '
          'request per search.',
          'Context provider values are memoised so a toast or theme change no longer re-renders '
          'the entire component tree.',
          'gzip compression on API responses, and on static assets at the nginx layer, with '
          'immutable caching for hashed asset filenames and no caching for index.html.',
          'A health endpoint that verifies database reachability and returns 503 when the '
          'database is down, so a broken deployment cannot report itself as healthy.',
          'Graceful shutdown on SIGTERM: in-flight requests are drained and the database pool '
          'is closed instead of being severed mid-request.',
          'Centralised error mapping so expected failures surface as the correct status code '
          '(409 for a duplicate record, 413 for an oversized upload, 503 for an unconfigured '
          'integration) rather than a generic 500.']:
    bullet(x)

doc.add_page_break()

# ============================================================
# 6. ARCHITECTURE
# ============================================================
h1('6. System Architecture')
para('ClientHub follows a classic decoupled three-tier architecture: a React SPA client, a '
     'stateless Express REST + Socket.io API server, and a PostgreSQL data tier accessed via '
     'Prisma. Third-party services (Groq AI, payment gateway, ImageKit, SMTP) are integrated '
     'through the server only, so secrets never reach the browser.')
code_block([
 "                         +-------------------------------------------+",
 "                         |              CLIENT (Browser)             |",
 "                         |   React + Tailwind + Recharts + dnd-kit   |",
 "                         |   Axios (REST)   |   Socket.io-client     |",
 "                         +---------+------------------+--------------+",
 "                                   | HTTPS / JWT      | WebSocket",
 "                                   v                  v",
 "                         +-------------------------------------------+",
 "                         |           SERVER (Node + Express)         |",
 "                         |  Routes -> Middleware -> Controllers      |",
 "                         |             -> Services (business logic)  |",
 "                         |  Auth(JWT) | RBAC | Validation | Sockets  |",
 "                         |  Cron Jobs (node-cron)                    |",
 "                         +--+--------+--------+--------+--------+-----+",
 "                            |        |        |        |        |",
 "            +---------------+   +----+--+  +--+----+  +-+------+ +----+-----+",
 "            v                   v       v  v       v  v        v v          v",
 "      +-----------+      +-----------+ +--------+ +-----------+ +-----------+",
 "      | PostgreSQL|      | Groq AI | | Razor- | | ImageKit  | |   SMTP    |",
 "      |  (Prisma) |      |   API     | | pay    | |  (logos)  | | (email)   |",
 "      +-----------+      +-----------+ +--------+ +-----------+ +-----------+",
])
para('Request flow: the browser sends a JWT-authenticated REST call → Express middleware '
     'validates the token and the user\'s role → the controller handles request/response only '
     '→ it delegates real work to a service → the service reads/writes the database via Prisma '
     'and calls external APIs as needed. Live updates are pushed back to the browser over Socket.io.',
     italic=True, color=GREY)

# ============================================================
# 7. DB SCHEMA
# ============================================================
h1('7. Database Schema (ER Description)')
para('The relational schema is normalized around nine core entities. Primary keys are id; FK '
     'denotes a foreign key relationship.')
kv_table(
    ['Entity', 'Key Fields', 'Relationships'],
    [
        ['User', 'id, name, email, passwordHash, role, createdAt', 'One User ↔ one Client account (for client role); assigned to many Tasks'],
        ['Client', 'id, name, companyName, contactEmail, contactPhone, notes, logoUrl, linkedUserId (FK)', 'Belongs to one User; has many Projects'],
        ['Project', 'id, clientId (FK), title, description, budget, startDate, deadline, status', 'Belongs to one Client; has many Tasks, Invoices, AISummaries'],
        ['Task', 'id, projectId (FK), title, description, assignedToUserId (FK), status, priority, dueDate, hoursLogged', 'Belongs to one Project; has many TaskComments'],
        ['TaskComment', 'id, taskId (FK), userId (FK), text, createdAt', 'Belongs to one Task and one User'],
        ['Invoice', 'id, projectId (FK), clientId (FK), lineItems (JSON), taxRate, totalAmount, status, dueDate, createdAt', 'Belongs to one Project & Client; has many Payments'],
        ['Payment', 'id, invoiceId (FK), amount, gatewayTransactionId, status, paidAt', 'Belongs to one Invoice'],
        ['Notification', 'id, userId (FK), message, type, isRead, createdAt', 'Belongs to one User'],
        ['AISummary', 'id, projectId (FK), summaryText, riskFlag, generatedAt', 'Belongs to one Project'],
        ['PasswordResetToken', 'id, tokenHash (SHA-256, unique), userId (FK), expiresAt, usedAt, createdAt', 'Belongs to one User; single-use, expires in 1 hour, redeemed atomically in a transaction'],
    ],
    widths=[1.1, 3.0, 2.4]
)

doc.add_page_break()

# ============================================================
# 8. TECH STACK
# ============================================================
h1('8. Technology Stack & Justification')
kv_table(
    ['Layer', 'Technology', 'Why it was chosen'],
    [
        ['Frontend', 'React + Tailwind CSS', 'Component-driven UI with fast, consistent utility-first styling.'],
        ['Charts', 'Recharts', 'Declarative, React-native charting for the analytics dashboards.'],
        ['Drag & Drop', '@dnd-kit', 'Modern, accessible drag-and-drop for the Kanban board.'],
        ['HTTP / Routing', 'Axios + React Router', 'Clean API layer and client-side route-based navigation.'],
        ['Backend', 'Node.js + Express', 'Lightweight, widely-adopted REST framework with a huge ecosystem.'],
        ['ORM', 'Prisma', 'Type-safe queries, easy migrations, readable schema.'],
        ['Database', 'PostgreSQL', 'Reliable relational DB ideal for related business data.'],
        ['Auth', 'JWT + bcrypt', 'Stateless scalable sessions; industry-standard password hashing.'],
        ['Real-time', 'Socket.io', 'Simple, reliable bi-directional events for live notifications.'],
        ['Payments', 'Razorpay (test mode)', 'Real gateway integration with HMAC-verified webhooks.'],
        ['AI', 'Groq API (Llama 3.3, OpenAI-compatible)', 'High-quality natural-language summaries and reasoning.'],
        ['Email', 'Nodemailer', 'Password resets and email fallback notifications.'],
        ['Media', 'ImageKit', 'Managed image hosting for client logos.'],
        ['PDF', 'pdfkit / puppeteer', 'Programmatic invoice PDF generation.'],
        ['Jobs', 'node-cron', 'Scheduled weekly AI summary generation.'],
        ['Validation', 'Zod / Joi', 'Schema-based request validation with clear errors.'],
        ['Logging', 'Winston / Morgan', 'Structured request and error logging.'],
        ['Testing', 'Vitest + Testing Library', 'Backend unit tests and frontend component/unit tests.'],
        ['Code Quality', 'ESLint + Prettier', 'Flat-config linting and consistent auto-formatting on both apps.'],
        ['CI', 'GitHub Actions', 'Lint + test + build run automatically on every push and pull request.'],
        ['Containerization', 'Docker + docker-compose', 'One-command local stack (Postgres + API + frontend).'],
        ['Deploy', 'Vercel + Render/Railway + Neon/Supabase', 'Free-tier friendly, fast CI/CD, managed Postgres.'],
    ],
    widths=[1.1, 2.1, 3.3]
)

doc.add_page_break()

# ============================================================
# 9. MODULE-WISE
# ============================================================
h1('9. Module-wise Explanation')

modules = [
    ('Module 1 — Authentication & Access Control',
     'Users sign up and log in with email and password. Passwords are never stored in plain text — '
     'they are hashed with bcrypt before saving. On successful login the server issues a signed JWT '
     'that the browser sends on every subsequent request. Middleware verifies the token and checks the '
     'user\'s role before allowing access, so a Client can never reach Admin-only endpoints. '
     'Forgot-password sends an emailed link containing a 32-byte random token; only its SHA-256 hash '
     'is stored in the PasswordResetToken table, so a database leak yields no usable links. Tokens are '
     'single-use, expire in one hour, and the password change plus token invalidation run in a single '
     'database transaction so they either both succeed or neither does. A daily cron job purges spent rows.'),
    ('Module 2 — Client Management',
     'The Admin manages a directory of clients, each with contact details, notes, and a logo uploaded to '
     'ImageKit. Every client record is linked to a client-role user account so the client can log into '
     'their own portal. The list supports search and filtering, and each client has a profile page listing '
     'all of their projects.'),
    ('Module 3 — Project Management',
     'Projects are created under a client with a budget, timeline, and status. Each project has a detail page '
     'showing its tasks, assigned team, and budget usage. A progress percentage is computed automatically from '
     'the ratio of completed tasks to total tasks, so progress is always accurate and never manually maintained.'),
    ('Module 4 — Task Management (Kanban & My Tasks)',
     'Tasks live inside projects and are visualized on a four-column Kanban board. Team members drag tasks '
     'between columns to change status (powered by @dnd-kit), discuss work in a per-task comment thread, and '
     'log hours. Every status change is recorded in an activity log for accountability. A dedicated My Tasks '
     'view aggregates every task assigned to the logged-in user across all projects, sorted by due date with '
     'overdue dates highlighted in red, so team members always know what needs attention first.'),
    ('Module 5 — Invoicing & Payments',
     'The invoice builder supports multiple line items with automatic tax and total calculation, and can '
     'auto-suggest line items from logged hours. Invoices export to PDF. Clients pay online through a payment '
     'gateway in test mode; a signature-verified webhook then automatically marks the invoice Paid, keeping '
     'billing state trustworthy and tamper-resistant.'),
    ('Module 6 — Real-Time Notifications',
     'Socket.io pushes live events to the right users the moment they happen — a task assignment, a new comment, '
     'a paid invoice, or a status change. A notification bell shows the unread count, and high-importance events '
     'also send an email fallback so nothing is missed.'),
    ('Module 7 — Analytics Dashboards',
     'The Admin dashboard visualizes revenue, pending payments, project status counts, team workload, and top '
     'clients. The Client dashboard shows their own project progress, upcoming deadlines, invoice status, and a '
     'recent activity feed. All charts are rendered with Recharts from live data.'),
    ('Module 8 — AI Progress Summary',
     'See the dedicated section below — this is the project\'s core differentiator.'),
]
for title, body in modules:
    h2(title)
    para(body)

doc.add_page_break()

# ============================================================
# 10. AI FEATURE
# ============================================================
h1('10. The AI Progress Summary Feature (Core Differentiator)')
para('This feature converts raw project data into client-ready communication automatically, and is the '
     'single biggest differentiator of ClientHub.')

h2('How it works, step by step')
numbered('Aggregation: aiSummaryService.js collects a project\'s recent activity — tasks completed this week, '
         'tasks still pending, total hours logged, days remaining until the deadline, and recent comments or '
         'flagged blockers.')
numbered('Prompt construction: that structured data is formatted into a clear, well-scoped prompt that asks the '
         'model to write a friendly, professional progress summary and to classify project risk.')
numbered('AI call: the prompt is sent to the Groq API via the Groq OpenAI-compatible REST API (via fetch).')
numbered('Response: the model returns (a) a natural-language summary of what was done, what is pending, and likely '
         'blockers, and (b) a risk flag — one of on-track, at-risk, or delayed.')
numbered('Persistence: the summary text and risk flag are saved to the AISummary table, preserving full history.')
numbered('Display: the latest summary appears on the Client dashboard with past summaries listed beneath it. '
         'Admins also have a "Generate Summary Now" button for on-demand generation.')
numbered('Automation: a weekly node-cron job runs the summary for every active project and can email the client.')

h2('Graceful failure handling')
para('The application must never crash because of an external AI call. If the Groq API is unreachable or errors, '
     'the service falls back to a deterministic template summary (for example, "X of Y tasks completed this week. '
     'Deadline in Z days."), sets the risk flag to "unknown", logs the error with Winston, and the dashboard shows '
     'a subtle "AI summary unavailable" badge. This makes the feature resilient and production-safe.')

# ============================================================
# 11. SECURITY
# ============================================================
h1('11. Security Considerations')
for x in [('Password hashing: ', 'bcrypt with salt (10 rounds); plain passwords are never stored.'),
          ('Secure password reset: ', 'a 32-byte random token is emailed; only its SHA-256 hash is stored in the '
                                      'PasswordResetToken table, so a database leak yields no usable reset links. '
                                      'Tokens are single-use (usedAt timestamp), expire in one hour, and the '
                                      'password update plus token invalidation run in a single database transaction.'),
          ('Stateless auth: ', 'signed JWTs with expiry, verified on every protected request. The '
                               'verification pins the HS256 algorithm, so a token signed with a '
                               'different scheme is rejected rather than accepted.'),
          ('Live revocation: ', 'the user record is re-read from the database on each request, so a '
                                'deleted or disabled account loses access immediately instead of '
                                'staying valid until the token expires.'),
          ('Role-based access: ', 'enforced on the backend (route guards) AND the frontend (protected routes).'),
          ('No privilege escalation: ', 'the public signup endpoint hard-codes the team_member role '
                                        'server-side; accepting a role from the request body would '
                                        'let anyone register themselves as an administrator.'),
          ('Tenant isolation: ', 'clients can only ever query their own data — never another client\'s. '
                                 'Each controller resolves the caller\'s own client record from their '
                                 'user id and filters by it, so a valid token still yields 403 when '
                                 'reaching for someone else\'s project, invoice or task.'),
          ('Webhook verification: ', 'payment webhooks are checked with an HMAC over the RAW request '
                                     'body before any event is trusted, which is why that route is '
                                     'mounted before the JSON body parser. Deliveries are also handled '
                                     'idempotently because gateways retry.'),
          ('Input validation: ', 'Zod schemas validate every request body, replacing it with the '
                                 'parsed and typed result, with clear per-field error messages.'),
          ('Rate limiting & CORS: ', 'authentication endpoints allow ten attempts per fifteen minutes '
                                     'per IP; CORS is restricted to a known origin and never falls '
                                     'back to a wildcard while credentials are enabled.'),
          ('Security headers: ', 'Helmet sets protections against XSS, clickjacking and MIME sniffing '
                                 'on the API; nginx adds the equivalent headers for the static frontend.'),
          ('Upload safety: ', 'logo uploads are held in memory, capped at 2 MB, and restricted to '
                              'image MIME types; violations return 413 or 400 rather than a 500.'),
          ('Error hygiene: ', 'stack traces and internal messages are never returned to clients in '
                              'production — one central handler decides exactly what is exposed.'),
          ('Secret management: ', 'all secrets live in environment variables; .env is never committed, '
                                  'and the server refuses to boot in production with a weak JWT secret.')]:
    bullet(x[1], x[0])

doc.add_page_break()

# ============================================================
# 12. FOLDER STRUCTURE
# ============================================================
h1('12. Final Project Folder Structure')
para('The project is organized as two clean applications — client (React frontend) and server '
     '(Express backend) — plus a docs package. Each responsibility lives in its own file; no giant '
     'single-file routes or components. The structure is self-explanatory so a reviewer can understand '
     'each folder by its name alone.')
code_block([
 "clienthub/",
 "|-- client/                          (React frontend)",
 "|   |-- src/",
 "|   |   |-- api/                     axiosClient.js, authApi.js, clientApi.js,",
 "|   |   |                            projectApi.js, taskApi.js, invoiceApi.js,",
 "|   |   |                            paymentApi.js, aiApi.js, notificationApi.js,",
 "|   |   |                            analyticsApi.js",
 "|   |   |-- components/",
 "|   |   |   |-- common/              Button, InputField, Modal, Loader, Badge,",
 "|   |   |   |                        EmptyState, Pagination, Logo, Toaster,",
 "|   |   |   |                        ErrorBoundary",
 "|   |   |   |-- layout/              Navbar, Sidebar, ProtectedRoute,",
 "|   |   |   |                        DashboardLayout",
 "|   |   |   |-- kanban/              KanbanBoard, TaskColumn, TaskCard, TaskModal",
 "|   |   |   |-- invoices/            InvoiceTable, InvoiceForm",
 "|   |   |   |-- dashboard/           StatsCard, RevenueChart, ProgressChart",
 "|   |   |   |-- notifications/       NotificationBell, NotificationList",
 "|   |   |-- pages/                   Login, Signup, ForgotPassword,",
 "|   |   |                            ResetPassword, AdminDashboard,",
 "|   |   |                            ClientDashboard, ClientsPage, ProjectsPage,",
 "|   |   |                            ProjectDetailPage, InvoicesPage, NotFound",
 "|   |   |-- context/                 contexts.js, AuthContext.jsx,",
 "|   |   |                            SocketContext.jsx, ThemeContext.jsx,",
 "|   |   |                            ToastContext.jsx",
 "|   |   |-- hooks/                   useAuth.js, useFetch.js, useDebounce.js,",
 "|   |   |                            useSocket.js, useTheme.js, useToast.js",
 "|   |   |-- utils/                   formatCurrency.js, dateHelpers.js,",
 "|   |   |                            roleGuards.js  (+ *.test.js unit tests)",
 "|   |   |-- test/                    setup.js  (Vitest + Testing Library)",
 "|   |-- nginx.conf                   (SPA fallback, gzip, cache headers)",
 "|   |   |-- routes/                  AppRoutes.jsx",
 "|   |   |-- App.jsx",
 "|   |   |-- main.jsx",
 "|   |-- Dockerfile                   (multi-stage build served by nginx)",
 "|   |-- eslint.config.js             (flat ESLint config)",
 "|   |-- vite.config.js               (build chunks + Vitest config)",
 "|   |-- .env.example",
 "|   |-- package.json",
 "|",
 "|-- server/                          (Node/Express backend)",
 "|   |-- src/",
 "|   |   |-- config/                  db.js, imagekit.js, razorpay.js,",
 "|   |   |                            groq.js",
 "|   |   |-- controllers/             authController.js, clientController.js,",
 "|   |   |                            projectController.js, taskController.js,",
 "|   |   |                            invoiceController.js, paymentController.js,",
 "|   |   |                            aiController.js, notificationController.js",
 "|   |   |-- routes/                  (one file per resource, matching controllers)",
 "|   |   |-- middleware/              authMiddleware.js, roleMiddleware.js,",
 "|   |   |                            errorHandler.js, rateLimiter.js,",
 "|   |   |                            validateRequest.js",
 "|   |   |-- services/                aiSummaryService.js, invoiceService.js,",
 "|   |   |                            paymentService.js, emailService.js,",
 "|   |   |                            notificationService.js",
 "|   |   |-- jobs/                    cronJobs.js",
 "|   |   |-- sockets/                 socketHandler.js",
 "|   |   |-- prisma/                  schema.prisma, migrations/, seed.js",
 "|   |   |-- utils/                   generatePDF.js, calculateProgress.js,",
 "|   |   |                            webhookVerify.js, logger.js",
 "|   |   |-- tests/                   calculations.test.js, auth.test.js,",
 "|   |   |                            aiSummary.test.js, security.test.js,",
 "|   |   |                            errorHandler.test.js  (Vitest)",
 "|   |   |-- app.js                   (Express app + middleware wiring)",
 "|   |   |-- server.js                (entry point: server + sockets)",
 "|   |-- Dockerfile                   (Node API image, runs migrations on boot)",
 "|   |-- eslint.config.js             (flat ESLint config)",
 "|   |-- .env.example",
 "|   |-- package.json",
 "|",
 "|-- .github/workflows/ci.yml         (GitHub Actions: lint + test + build)",
 "|-- docker-compose.yml               (Postgres + API + frontend, one command)",
 "|-- .prettierrc.json / .prettierignore",
 "|-- docs/                            ClientHub_Project_Report.docx,",
 "|                                    ClientHub_Presentation.pptx,",
 "|                                    ClientHub.postman_collection.json,",
 "|                                    ClientHub_TODO_Guide.docx",
 "|-- README.md / README.txt",
 "|-- prompt.txt, run.txt, resume.txt, WEBSITE_GUIDE.txt",
 "|-- best free way to live this project for free.txt",
 "|-- LICENSE",
 "|-- .gitignore",
])
para('Design rule: controllers handle only request/response; all business logic (calculations, AI '
     'calls, payment logic) lives in services/ and is imported into controllers. React pages are thin '
     'compositions of reusable components.', italic=True, color=GREY)

doc.add_page_break()

# ============================================================
# 13. SEED & DEMO
# ============================================================
h1('13. Setup, Seed Data & Demo Credentials')
para('A seed script populates the database with realistic demo data — several clients, projects, tasks '
     'in different statuses, sample invoices (some paid, some pending), and a sample AI summary — so the '
     'live demo looks populated, not empty. The seed is self-guarding: it skips itself when the database '
     'already contains users, so restarting the stack never destroys existing data (set FORCE_SEED=true '
     'to deliberately wipe and reseed). Fixed demo logins exist for every role; they are kept out of the '
     'login screen so the demo reads as a real product rather than a tutorial.')
kv_table(
    ['Role', 'Email', 'Password'],
    [['Admin (Agency Owner)', 'admin@demo.com', 'Demo@1234'],
     ['Team Member', 'team@demo.com', 'Demo@1234'],
     ['Client', 'client@demo.com', 'Demo@1234']],
    widths=[2.2, 2.2, 2.0]
)
para('Note: these are demo-only credentials for a test deployment with non-sensitive seeded data.',
     italic=True, color=LIGHT, size=10)

# ============================================================
# 14. TESTING & API DOCS
# ============================================================
h1('14. Testing, Code Quality & CI/CD')
bullet(' 44 automated tests run with Vitest and all pass. On the backend (29): the auth controller '
       '(signup conflicts, wrong-password handling, token issuance), the progress and invoice-total '
       'calculations, the AI summary service\'s fallback path, the RBAC middleware, the central error '
       'handler, and the payment webhook signature verification. On the frontend (15): the currency, '
       'date and role-guard helpers, the debounce hook, and the Badge component.', 'Automated tests:')
bullet(' the security-critical tests are deliberately adversarial rather than happy-path — the '
       'webhook suite asserts that a signature made with the wrong secret is rejected, that a '
       'tampered body invalidates a previously valid signature, and that a wrong-length signature '
       'is refused without throwing. The error-handler suite asserts that no internal message or '
       'stack detail reaches the client on a 500 in production.', 'Testing approach:')
bullet(' both the client and server are linted with ESLint (flat config) and auto-formatted '
       'with Prettier. Both currently report zero errors and zero warnings.', 'Linting & formatting:')
bullet(' a GitHub Actions workflow (.github/workflows/ci.yml) automatically runs lint, tests, '
       'and a production build for both apps on every push and pull request to main.', 'Continuous Integration:')
bullet(' route-level code splitting with React.lazy, plus a dynamic import for the Socket.io '
       'client and Vite manualChunks for the remaining vendor libraries. The initial payload fell '
       'from roughly 748 KB to 233 KB (about 79 KB gzipped) because the charting and drag-and-drop '
       'libraries now load only on the pages that need them. A chunk-size budget in the Vite '
       'config warns if that regresses.', 'Build optimization:')
bullet(' a Postman collection (docs/ClientHub.postman_collection.json) covering all endpoints, '
       'plus a health-check endpoint at /api/health that also reports database reachability and '
       'process uptime, and returns 503 when the database is unreachable.', 'API docs:')
bullet(' structured request and error logging with Winston/Morgan instead of scattered console.log '
       'calls, with verbose logging in development and combined-format access logs in production.', 'Logging:')
bullet(' page-based pagination and search/filter on all large list views (clients, tasks, invoices), '
       'with the search input debounced so typing does not issue a request per keystroke.', 'Pagination:')
bullet(' the environment is validated at boot. Missing required variables abort the start with a '
       'message naming them; a weak JWT secret is fatal in production; and missing optional keys '
       'are logged so it is always obvious why an integration is idle.', 'Configuration safety:')

# ============================================================
# 15. DEPLOYMENT
# ============================================================
h1('15. Deployment Architecture')
kv_table(
    ['Component', 'Platform', 'Notes'],
    [['Frontend (React)', 'Vercel', 'Automatic builds from Git; global CDN.'],
     ['Backend (Express)', 'Render / Railway', 'Hosts API + Socket.io; environment secrets configured in dashboard.'],
     ['Database (Postgres)', 'Neon / Supabase', 'Managed free-tier PostgreSQL with connection pooling.'],
     ['Media', 'ImageKit', 'Client logo storage.'],
     ['AI / Payments / Email', 'Groq API / Razorpay / SMTP', 'Accessed server-side using environment secrets.'],
     ['Local / self-host', 'Docker + docker-compose', 'One command spins up Postgres, API, and frontend together.'],
     ['CI pipeline', 'GitHub Actions', 'Runs lint, tests, and build on every push/PR before deploy.']],
    widths=[1.8, 1.8, 2.8]
)

doc.add_page_break()

# ============================================================
# 16. SCREENSHOTS
# ============================================================
h1('16. Screenshots (Placeholders)')
para('Insert screenshots of the running application here before final submission:')
for s in ['[ Screenshot 1: Login page (light and dark mode side by side) ]',
          '[ Screenshot 2: Admin analytics dashboard (revenue & project charts) ]',
          '[ Screenshot 3: Kanban board with drag-and-drop tasks ]',
          '[ Screenshot 4: Invoice builder and PDF preview ]',
          '[ Screenshot 5: Client portal with the AI progress summary ]',
          '[ Screenshot 6: Real-time notification bell with unread count ]']:
    p = para(s, align=WD_ALIGN_PARAGRAPH.CENTER, color=LIGHT)
    p.paragraph_format.space_after = Pt(10)

# ============================================================
# 17. CHALLENGES
# ============================================================
h1('17. Challenges Faced & Solutions')
kv_table(
    ['Challenge', 'Solution'],
    [
        ['Keeping controllers clean as features grew', 'Separated all business logic into a services/ layer, leaving controllers as thin request/response handlers.'],
        ['Trusting payment status from the client', 'Relied on signature-verified gateway webhooks to update invoice status server-side, never the browser.'],
        ['Password reset tokens were stored in memory', 'An in-memory Map breaks the moment the server restarts or runs as more than one instance — the user\'s reset link silently stops working. Moved tokens to a dedicated PasswordResetToken table, storing only the SHA-256 hash, with single-use enforcement and atomic redemption in a database transaction.'],
        ['AI API could fail or be slow', 'Added a deterministic template fallback, Winston error logging, an "unknown" risk flag, and an "AI unavailable" badge.'],
        ['Enforcing tenant isolation', 'Scoped every client-facing query by the authenticated user so clients can only read their own data.'],
        ['Real-time updates to the right users only', 'Used Socket.io rooms keyed by user id to target events precisely.'],
        ['Accurate project progress', 'Computed progress from completed-vs-total tasks instead of manual entry.'],
        ['A drag also fired a click, so dropping a Kanban card opened its detail modal',
         'Tracked whether the pointer gesture became a drag and swallowed the click that follows it on pointer-up.'],
        ['Real-time notifications worked locally but would have failed once deployed',
         'The Socket.io server allowed origin "*" together with credentials — a combination browsers reject. Pinned it to the configured client origin, matching the REST CORS policy.'],
        ['Overdue invoices became unpayable',
         'The daily job moves a past-due invoice from pending to overdue, but the Pay button only appeared for pending ones. Widened the condition to any unpaid invoice.'],
        ['Every visitor downloaded the charting library just to see the login page',
         'Introduced route-level code splitting with React.lazy and a dynamic import for the socket client, cutting the initial payload by about 69%.'],
        ['Typing in the client search fired one API request per keystroke',
         'Added a debounce hook so only the settled value drives the query, with unit tests proving rapid typing collapses to a single update.'],
        ['Scheduled jobs fired at the wrong local time once hosted',
         'Hosts run in UTC, so "Monday 09:00" arrived mid-afternoon IST. Pinned both cron jobs to an explicit, configurable timezone.'],
        ['Expected failures looked like server crashes',
         'Mapped Prisma and upload errors centrally, so a duplicate email returns 409 and an oversized file returns 413 instead of an opaque 500.'],
        ['Dark mode flashed the light theme on every page load',
         'Moved the theme decision into a tiny inline script that runs before React mounts and first paint.'],
        ['The application could not be navigated on a phone',
         'The sidebar was hidden below tablet width with nothing to replace it. Added a drawer in the navbar sharing one link definition with the desktop rail so the two cannot drift apart.'],
    ],
    widths=[2.6, 3.8]
)

# ============================================================
# 18. FUTURE SCOPE
# ============================================================
h1('18. Future Scope')
for x in ['Subscription billing and recurring invoices for retainer clients.',
          'A mobile app (React Native) sharing the same API.',
          'AI-assisted task estimation and automatic deadline-risk prediction.',
          'Client e-signature on proposals and contracts.',
          'Multi-currency support and region-specific tax rules.',
          'Granular team permissions and audit trails.',
          'In-app chat between agency and client.']:
    bullet(x)

# ============================================================
# 19. CONCLUSION
# ============================================================
h1('19. Conclusion')
para('ClientHub demonstrates the ability to design and build a complete, real-world SaaS product '
     'end-to-end: secure multi-role authentication, relational data modeling, a polished React UI, '
     'real-time communication, third-party payment and email integrations, automated background jobs, '
     'and a genuinely useful AI feature with production-grade failure handling. Its clean, well-commented, '
     'professionally organized codebase makes it both a strong academic submission and a credible portfolio '
     'project that a recruiter can open, understand, and trust.')

# ============================================================
# 20. RESUME SUMMARY
# ============================================================
h1('20. Resume Summary')
para('Suggested resume bullet:', bold=True)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
r = p.add_run('Built ClientHub, an AI-powered agency management SaaS (React, Node/Express, Prisma, '
              'PostgreSQL, Socket.io) with multi-role access control, drag-and-drop Kanban, invoicing '
              'with webhook-verified online payments, real-time notifications, and a Groq-powered weekly '
              'progress-report engine with risk flagging and graceful fallback — tested with Vitest, '
              'linted with ESLint/Prettier, containerized with Docker, CI via GitHub Actions, and '
              'deployed full-stack on Vercel, Render, and Neon.')
r.font.italic = True
r.font.size = Pt(11)
r.font.color.rgb = NAVY

doc.add_paragraph()
para('— Prepared by Suryakanta Pradhan (MCA)', align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, color=GREY)

doc.save('clienthub/docs/ClientHub_Project_Report.docx')
print('DOCX saved OK')
